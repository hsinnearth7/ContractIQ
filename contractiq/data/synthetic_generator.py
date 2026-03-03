"""Generate synthetic supplier contracts (PDF/DOCX) for demo and testing."""

import random
from datetime import datetime, timedelta
from pathlib import Path

from fpdf import FPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from contractiq.data.clause_library import CLAUSE_LIBRARY, CLAUSE_NAMES
from contractiq.data.contract_templates import TEMPLATES


# --- Supplier pool ---
SUPPLIERS = [
    {"name": "Acme Technologies Inc.", "address": "123 Innovation Blvd, San Jose, CA 95134"},
    {"name": "GlobalLogistics Corp.", "address": "456 Harbor Drive, Long Beach, CA 90802"},
    {"name": "PrimeData Solutions", "address": "789 Analytics Way, Austin, TX 78701"},
    {"name": "NexGen Cloud Services", "address": "321 Cloud Ave, Seattle, WA 98101"},
    {"name": "Pinnacle Manufacturing Ltd.", "address": "555 Industrial Park, Detroit, MI 48201"},
    {"name": "Vertex Consulting Group", "address": "888 Strategy Lane, Chicago, IL 60601"},
    {"name": "BlueOcean IT Services", "address": "222 Tech Center, Raleigh, NC 27601"},
    {"name": "SilverLine Distributors", "address": "444 Commerce St, Dallas, TX 75201"},
    {"name": "EcoSupply Partners", "address": "777 Green Blvd, Portland, OR 97201"},
    {"name": "CyberShield Security", "address": "999 Defense Circle, Arlington, VA 22201"},
]

BUYER = {
    "name": "TechCorp International",
    "address": "100 Enterprise Way, San Francisco, CA 94105",
}

SIGNATORIES = [
    ("John Smith", "VP of Procurement"),
    ("Sarah Chen", "Director of Operations"),
    ("Michael Brown", "Chief Procurement Officer"),
    ("Emily Davis", "Head of Vendor Management"),
    ("Robert Wilson", "Senior Contract Manager"),
]

SERVICE_DESCRIPTIONS = {
    "MSA": [
        "enterprise software development, maintenance, and support services",
        "cloud infrastructure management and DevOps consulting services",
        "data analytics platform implementation and ongoing optimization",
        "supply chain management software licensing and customization",
        "cybersecurity assessment, monitoring, and incident response services",
    ],
    "PO": [
        "500 units of Model X-200 server hardware with 3-year warranty",
        "enterprise software licenses for 1,000 users including annual maintenance",
        "industrial automation components including PLCs and HMI panels",
        "networking equipment: 50 switches, 20 routers, 100 access points",
        "office furniture and ergonomic workstations for new facility",
    ],
    "NDA": [
        "potential partnership in AI/ML-powered supply chain optimization",
        "evaluation of proprietary logistics routing algorithm",
        "joint development of blockchain-based procurement platform",
        "assessment of quantum computing applications in supply chain",
    ],
    "SLA": [
        "managed cloud hosting and application support services",
        "24/7 IT helpdesk and infrastructure monitoring services",
        "enterprise ERP system hosting and administration",
        "disaster recovery and business continuity services",
    ],
}


def _random_date(start_year: int = 2023, end_year: int = 2025) -> str:
    start = datetime(start_year, 1, 1)
    end = datetime(end_year, 12, 31)
    delta = end - start
    rand_date = start + timedelta(days=random.randint(0, delta.days))
    return rand_date.strftime("%B %d, %Y")


def _pick_clauses(
    contract_type: str,
    include_all: bool = True,
    omit_clauses: list[str] | None = None,
) -> list[tuple[str, str]]:
    """Select clauses for a contract, optionally omitting some for compliance testing."""
    omit = set(omit_clauses or [])
    available = [c for c in CLAUSE_NAMES if c not in omit]

    if contract_type == "NDA":
        # NDAs have fewer clauses
        preferred = ["governing_law", "dispute_resolution", "entire_agreement", "notice"]
        available = [c for c in preferred if c not in omit]

    clauses = []
    for name in available:
        variants = CLAUSE_LIBRARY[name]
        text = random.choice(variants)
        title = name.replace("_", " ").title()
        clauses.append((title, text))

    return clauses


def _build_context(
    contract_type: str,
    supplier: dict,
    effective_date: str,
    idx: int,
) -> dict:
    """Build template rendering context."""
    buyer_sig, buyer_title = random.choice(SIGNATORIES)
    sup_sig, sup_title = random.choice(SIGNATORIES)

    ctx = {
        "agreement_number": f"CIQ-{contract_type}-{2024}-{idx:03d}",
        "effective_date": effective_date,
        "buyer_name": BUYER["name"],
        "buyer_address": BUYER["address"],
        "buyer_signatory": buyer_sig,
        "buyer_title": buyer_title,
        "supplier_name": supplier["name"],
        "supplier_address": supplier["address"],
        "supplier_signatory": sup_sig,
        "supplier_title": sup_title,
        "service_description": random.choice(SERVICE_DESCRIPTIONS[contract_type]),
    }

    if contract_type == "MSA":
        ctx.update({
            "term_years": random.choice([1, 2, 3]),
            "notice_days": random.choice([30, 60, 90]),
            "contract_value": f"{random.randint(100, 5000) * 1000:,}",
        })
    elif contract_type == "PO":
        unit_price = random.randint(50, 5000)
        quantity = random.randint(10, 1000)
        ctx.update({
            "delivery_date": _random_date(2024, 2025),
            "quantity": quantity,
            "unit_price": f"{unit_price:,}",
            "contract_value": f"{unit_price * quantity:,}",
            "shipping_terms": random.choice(["FOB Destination", "FOB Origin", "CIF"]),
            "acceptance_days": random.choice([5, 10, 15]),
        })
    elif contract_type == "NDA":
        ctx.update({
            "term_years": random.choice([2, 3, 5]),
            "survival_years": random.choice([3, 5, 7]),
            "return_days": random.choice([10, 15, 30]),
        })
    elif contract_type == "SLA":
        ctx.update({
            "associated_contract": f"CIQ-MSA-2024-{random.randint(1,10):03d}",
            "uptime_pct": random.choice(["99.5", "99.9", "99.95"]),
            "measurement_period": random.choice(["monthly", "quarterly"]),
            "p1_response": "15 minutes",
            "p1_resolution": "4 hours",
            "p2_response": "1 hour",
            "p2_resolution": "8 hours",
            "p3_response": "4 hours",
            "p3_resolution": "2 business days",
            "p4_response": "1 business day",
            "p4_resolution": "5 business days",
            "maintenance_window": "Saturday 2:00 AM - 6:00 AM ET",
            "maintenance_notice": random.choice(["48", "72"]),
            "credit_tier1_pct": "99.9",
            "credit_tier1_amount": "5",
            "credit_tier2_pct": "99.5",
            "credit_tier2_amount": "10",
            "credit_tier3_pct": "99.0",
            "credit_tier3_amount": "20",
            "max_credit_pct": "30",
            "contract_value": f"{random.randint(5, 50) * 1000:,}",
            "billing_period": random.choice(["month", "quarter"]),
            "report_days": random.choice(["5", "10"]),
            "escalation_l1": "support@supplier.com",
            "escalation_l2": "service.manager@supplier.com",
            "escalation_l3": "vp.operations@supplier.com",
        })

    return ctx


def _render_contract(contract_type: str, context: dict, clauses: list[tuple[str, str]]) -> str:
    """Render a contract using Jinja2 template."""
    template = TEMPLATES[contract_type]
    context["clauses"] = clauses
    return template.render(**context)


def _save_as_pdf(text: str, filepath: Path) -> None:
    """Save contract text as PDF."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=9)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue

        # Encode to latin-1 safe text for FPDF
        safe_line = line.encode("latin-1", errors="replace").decode("latin-1")

        # Detect headings
        mc_kwargs = {"w": 0, "new_x": "LMARGIN", "new_y": "NEXT"}
        if safe_line.isupper() and len(safe_line) < 80:
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(h=5, text=safe_line, **mc_kwargs)
            pdf.set_font("Helvetica", size=9)
        elif safe_line.startswith("ARTICLE") or safe_line.startswith("BETWEEN") or safe_line.startswith("AND:"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(h=5, text=safe_line, **mc_kwargs)
            pdf.set_font("Helvetica", size=9)
        else:
            pdf.multi_cell(h=4, text=safe_line, **mc_kwargs)

    pdf.output(str(filepath))


def _save_as_docx(text: str, filepath: Path) -> None:
    """Save contract text as DOCX."""
    doc = DocxDocument()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.isupper() and len(line) < 80:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith("ARTICLE") or line.startswith("BETWEEN") or line.startswith("AND:"):
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(11)
        else:
            doc.add_paragraph(line)

    doc.save(str(filepath))


# Contracts that intentionally omit clauses for compliance testing
COMPLIANCE_TEST_OMISSIONS = {
    3: ["force_majeure", "data_protection"],
    7: ["indemnification", "limitation_of_liability", "insurance"],
    11: ["termination", "dispute_resolution"],
    15: ["confidentiality", "warranty"],
    19: ["force_majeure", "payment_terms", "governing_law"],
}


def generate_contracts(
    output_dir: str | Path,
    count: int = 20,
    seed: int = 42,
) -> list[Path]:
    """Generate synthetic contracts and save as PDF/DOCX.

    Args:
        output_dir: Directory to save generated contracts.
        count: Number of contracts to generate.
        seed: Random seed for reproducibility.

    Returns:
        List of generated file paths.
    """
    random.seed(seed)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    contract_types = ["MSA", "MSA", "PO", "PO", "NDA", "SLA"]  # weighted distribution
    generated_files: list[Path] = []

    for i in range(1, count + 1):
        ct = contract_types[(i - 1) % len(contract_types)]
        supplier = SUPPLIERS[(i - 1) % len(SUPPLIERS)]
        date = _random_date()

        # Build omission list for compliance testing
        omit = COMPLIANCE_TEST_OMISSIONS.get(i, [])

        clauses = _pick_clauses(ct, omit_clauses=omit)
        context = _build_context(ct, supplier, date, i)
        text = _render_contract(ct, context, clauses)

        # Alternate between PDF and DOCX
        supplier_slug = supplier["name"].split()[0].lower()
        if i % 3 == 0:
            filename = f"{ct.lower()}_{supplier_slug}_{i:03d}.docx"
            filepath = output_path / filename
            _save_as_docx(text, filepath)
        else:
            filename = f"{ct.lower()}_{supplier_slug}_{i:03d}.pdf"
            filepath = output_path / filename
            _save_as_pdf(text, filepath)

        generated_files.append(filepath)
        print(f"  Generated: {filename} ({ct}, {supplier['name']})"
              + (f" [OMITTED: {', '.join(omit)}]" if omit else ""))

    return generated_files


if __name__ == "__main__":
    from contractiq.config import get_settings
    settings = get_settings()
    print(f"Generating contracts to: {settings.sample_contracts_dir}")
    files = generate_contracts(settings.sample_contracts_dir, count=20)
    print(f"\nGenerated {len(files)} contracts.")
