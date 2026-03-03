"""DOCX parsing using python-docx."""

from pathlib import Path

from docx import Document as DocxDocument

from contractiq.parsing.models import ParsedDocument


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """Parse a DOCX file extracting text and tables."""
    file_path = Path(file_path)
    errors: list[str] = []

    try:
        doc = DocxDocument(str(file_path))
    except Exception as e:
        return ParsedDocument(
            file_name=file_path.name,
            file_path=str(file_path),
            file_type="docx",
            raw_text="",
            parse_errors=[f"Failed to open DOCX: {e}"],
        )

    # --- Extract paragraphs ---
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect heading styles
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip()
            prefix = "#" * int(level) if level.isdigit() else "##"
            paragraphs.append(f"{prefix} {text}")
        elif any(run.bold for run in para.runs if run.text.strip()):
            # Bold text treated as sub-heading
            paragraphs.append(f"**{text}**")
        else:
            paragraphs.append(text)

    raw_text = "\n\n".join(paragraphs)

    # --- Extract tables ---
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    # Approximate page count (DOCX doesn't store this natively)
    page_count = max(1, len(raw_text) // 3000)

    return ParsedDocument(
        file_name=file_path.name,
        file_path=str(file_path),
        file_type="docx",
        raw_text=raw_text,
        page_count=page_count,
        tables=tables,
        parse_errors=errors,
    )
